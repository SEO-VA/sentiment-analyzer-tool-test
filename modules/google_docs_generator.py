# modules/google_docs_generator.py

import html
from typing import List, Dict, Any, Optional
from io import BytesIO
from datetime import datetime

def generate_google_docs_files(sentences: List[Dict[str, Any]], results: List[Dict[str, Any]], 
                              webpage_data: Optional[Dict[str, Any]] = None) -> Dict[str, bytes]:
    """
    Generate files optimized for Google Docs import with full formatting preservation
    
    Returns multiple format options:
    - Rich Text Format (.rtf) - Best for Google Docs import with colors
    - Microsoft Word (.docx) - Alternative with full formatting (if python-docx available)
    - Enhanced HTML (.html) - Optimized for Google Docs import
    
    Args:
        sentences: List of sentence data
        results: Classification results  
        webpage_data: Optional webpage data
        
    Returns:
        Dict with format names as keys and file content as bytes
    """
    files = {}
    
    # Generate RTF (Rich Text Format) - Best for Google Docs color preservation
    files['rtf'] = _generate_rtf_content(sentences, results, webpage_data)
    
    # Generate enhanced HTML optimized for Google Docs import
    files['html'] = _generate_google_docs_html(sentences, results, webpage_data)
    
    # Generate Word document (DOCX) if possible
    try:
        files['docx'] = _generate_docx_content(sentences, results, webpage_data)
    except ImportError:
        # python-docx not available, skip DOCX generation
        pass
    except Exception:
        # Any other error in DOCX generation, skip it
        pass
    
    return files

def _generate_rtf_content(sentences: List[Dict[str, Any]], results: List[Dict[str, Any]], 
                         webpage_data: Optional[Dict[str, Any]] = None) -> bytes:
    """
    Generate RTF (Rich Text Format) with color highlighting
    RTF preserves colors and formatting when imported to Google Docs
    """
    
    # RTF color table - defines the highlight colors
    color_table = r"""{\colortbl;
\red173\green216\blue230;  
\red240\green128\blue128;  
\red144\green238\blue144;  
\red0\green0\blue0;        
}"""
    
    # RTF header
    rtf_content = [
        r"{\rtf1\ansi\deff0",
        r"{\fonttbl{\f0\froman Times New Roman;}}",
        color_table
    ]
    
    # Document title and metadata
    title = "Content Classification Results"
    if webpage_data and webpage_data.get('title'):
        title = f"Content Classification: {webpage_data['title']}"
    
    rtf_content.extend([
        r"\f0\fs24",  # Font and size
        r"{\b " + _rtf_escape(title) + r"}\par\par"
    ])
    
    # Add source information if available
    if webpage_data and webpage_data.get('success'):
        rtf_content.extend([
            r"{\b Source Information:}\par",
            f"Title: {_rtf_escape(webpage_data.get('title', ''))}\\par",
        ])
        if webpage_data.get('url'):
            rtf_content.append(f"URL: {_rtf_escape(webpage_data['url'])}\\par")
        rtf_content.append(r"\par")
    
    # Add statistics
    stats = _calculate_statistics(sentences, results)
    rtf_content.extend([
        r"{\b Classification Summary:}\par",
        f"Informational: {stats['info_pct']}% ({stats['info_count']} characters)\\par",
        f"Promotional: {stats['promo_pct']}% ({stats['promo_count']} characters)\\par", 
        f"Risk Warning: {stats['risk_pct']}% ({stats['risk_count']} characters)\\par",
        f"Total Items: {stats['total_items']}\\par\\par"
    ])
    
    # Add legend
    rtf_content.extend([
        r"{\b Legend:} ",
        r"{\highlight1 Informational} ",
        r"{\highlight2 Promotional} ",
        r"{\highlight3 Risk Warning}\par\par"
    ])
    
    # Add classified content with highlighting
    rtf_content.append(r"{\b Classified Content:}\par")
    
    # Check if we have webpage structure to preserve
    if webpage_data and webpage_data.get('structure'):
        # Use structure-aware rendering for RTF
        structured_content = _generate_rtf_with_structure(sentences, results, webpage_data)
        rtf_content.append(structured_content)
    else:
        # Use simple sentence-by-sentence rendering
        for result in results:
            idx = result["idx"]
            sentence = sentences[idx]["content"]
            
            if "spans" in result:
                # Handle phrase-level spans
                for span in result["spans"]:
                    text_part = sentence[span["start"]:span["end"]]
                    color_code = _get_rtf_color_code(span["label"])
                    escaped_text = _rtf_escape(text_part)
                    rtf_content.append(f"{{\\highlight{color_code} {escaped_text}}}")
            else:
                # Handle sentence-level classification
                color_code = _get_rtf_color_code(result["label"])
                escaped_text = _rtf_escape(sentence)
                rtf_content.append(f"{{\\highlight{color_code} {escaped_text}}} ")
    
    rtf_content.append("}")  # Close RTF document
    
    return "\n".join(rtf_content).encode('utf-8')

def _generate_google_docs_html(sentences: List[Dict[str, Any]], results: List[Dict[str, Any]], 
                              webpage_data: Optional[Dict[str, Any]] = None) -> bytes:
    """
    Generate HTML specifically optimized for Google Docs import
    Uses inline styles that Google Docs recognizes
    """
    
    title = "Content Classification Results"
    if webpage_data and webpage_data.get('title'):
        title = f"Content Classification: {webpage_data['title']}"
    
    stats = _calculate_statistics(sentences, results)
    
    html_parts = [
        '<!DOCTYPE html>',
        '<html>',
        '<head>',
        '<meta charset="UTF-8">',
        f'<title>{html.escape(title)}</title>',
        '<style>',
        'body { font-family: Arial, sans-serif; font-size: 11pt; line-height: 1.15; margin: 72pt; }',
        '.title { font-size: 18pt; font-weight: bold; margin-bottom: 12pt; }',
        '.subtitle { font-size: 14pt; font-weight: bold; margin-top: 18pt; margin-bottom: 6pt; }',
        '.stats { margin-bottom: 12pt; }',
        '.legend { margin-bottom: 12pt; font-weight: bold; }',
        '.info { background-color: #ADD8E6; }',      # Light blue
        '.promo { background-color: #F08080; }',     # Light coral  
        '.risk { background-color: #90EE90; }',      # Light green
        '.content { margin-top: 12pt; }',
        '</style>',
        '</head>',
        '<body>',
        
        f'<div class="title">{html.escape(title)}</div>'
    ]
    
    # Add source information
    if webpage_data and webpage_data.get('success'):
        html_parts.extend([
            '<div class="subtitle">Source Information</div>',
            f'<p><strong>Title:</strong> {html.escape(webpage_data.get("title", ""))}</p>'
        ])
        if webpage_data.get('url'):
            html_parts.append(f'<p><strong>URL:</strong> {html.escape(webpage_data["url"])}</p>')
    
    # Add statistics
    html_parts.extend([
        '<div class="subtitle">Classification Summary</div>',
        '<div class="stats">',
        f'<p><strong>Informational:</strong> {stats["info_pct"]}% ({stats["info_count"]:,} characters)</p>',
        f'<p><strong>Promotional:</strong> {stats["promo_pct"]}% ({stats["promo_count"]:,} characters)</p>',
        f'<p><strong>Risk Warning:</strong> {stats["risk_pct"]}% ({stats["risk_count"]:,} characters)</p>',
        f'<p><strong>Total Items:</strong> {stats["total_items"]}</p>',
        '</div>'
    ])
    
    # Add legend
    html_parts.extend([
        '<div class="legend">',
        'Legend: ',
        '<span class="info">Informational</span> ',
        '<span class="promo">Promotional</span> ',
        '<span class="risk">Risk Warning</span>',
        '</div>'
    ])
    
    # Add classified content
    html_parts.extend([
        '<div class="subtitle">Classified Content</div>',
        '<div class="content">'
    ])
    
    # Check if we have webpage structure to preserve
    if webpage_data and webpage_data.get('structure'):
        # Use structure-aware rendering for HTML
        structured_html = _generate_html_with_structure(sentences, results, webpage_data)
        html_parts.append(structured_html)
    else:
        # Use simple sentence-by-sentence rendering
        for result in results:
            idx = result["idx"]
            sentence = sentences[idx]["content"]
            
            if "spans" in result:
                # Handle phrase-level spans
                for span in result["spans"]:
                    text_part = sentence[span["start"]:span["end"]]
                    css_class = span["label"]  # 'info', 'promo', or 'risk'
                    escaped_text = html.escape(text_part)
                    html_parts.append(f'<span class="{css_class}">{escaped_text}</span>')
            else:
                # Handle sentence-level classification
                css_class = result["label"]
                escaped_text = html.escape(sentence)
                html_parts.append(f'<span class="{css_class}">{escaped_text}</span> ')
    
    html_parts.extend([
        '</div>',
        '</body>',
        '</html>'
    ])
    
    return '\n'.join(html_parts).encode('utf-8')

def _generate_docx_content(sentences: List[Dict[str, Any]], results: List[Dict[str, Any]], 
                          webpage_data: Optional[Dict[str, Any]] = None) -> bytes:
    """
    Generate Microsoft Word document with highlighting
    Requires python-docx package
    """
    from docx import Document
    from docx.shared import RGBColor
    from docx.enum.text import WD_COLOR_INDEX
    
    doc = Document()
    
    # Document title
    title = "Content Classification Results"
    if webpage_data and webpage_data.get('title'):
        title = f"Content Classification: {webpage_data['title']}"
    
    title_para = doc.add_heading(title, level=1)
    
    # Add source information
    if webpage_data and webpage_data.get('success'):
        doc.add_heading('Source Information', level=2)
        doc.add_paragraph(f"Title: {webpage_data.get('title', '')}")
        if webpage_data.get('url'):
            doc.add_paragraph(f"URL: {webpage_data['url']}")
    
    # Add statistics
    stats = _calculate_statistics(sentences, results)
    doc.add_heading('Classification Summary', level=2)
    stats_para = doc.add_paragraph()
    stats_para.add_run(f"Informational: {stats['info_pct']}% ({stats['info_count']:,} characters)\n")
    stats_para.add_run(f"Promotional: {stats['promo_pct']}% ({stats['promo_count']:,} characters)\n")
    stats_para.add_run(f"Risk Warning: {stats['risk_pct']}% ({stats['risk_count']:,} characters)\n")
    stats_para.add_run(f"Total Items: {stats['total_items']}")
    
    # Add legend
    legend_para = doc.add_paragraph("Legend: ")
    
    info_run = legend_para.add_run("Informational ")
    info_run.font.highlight_color = WD_COLOR_INDEX.TURQUOISE
    
    promo_run = legend_para.add_run("Promotional ")
    promo_run.font.highlight_color = WD_COLOR_INDEX.PINK
    
    risk_run = legend_para.add_run("Risk Warning")
    risk_run.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
    
    # Add classified content
    doc.add_heading('Classified Content', level=2)
    
    # Check if we have webpage structure to preserve
    if webpage_data and webpage_data.get('structure'):
        # Use structure-aware rendering for DOCX
        _generate_docx_with_structure(doc, sentences, results, webpage_data)
    else:
        # Use simple paragraph approach
        content_para = doc.add_paragraph()
        
        # Color mapping for Word highlighting
        color_map = {
            'info': WD_COLOR_INDEX.TURQUOISE,
            'promo': WD_COLOR_INDEX.PINK, 
            'risk': WD_COLOR_INDEX.BRIGHT_GREEN
        }
        
        for result in results:
            idx = result["idx"]
            sentence = sentences[idx]["content"]
            
            if "spans" in result:
                # Handle phrase-level spans
                for span in result["spans"]:
                    text_part = sentence[span["start"]:span["end"]]
                    run = content_para.add_run(text_part)
                    run.font.highlight_color = color_map[span["label"]]
            else:
                # Handle sentence-level classification
                run = content_para.add_run(sentence + " ")
                run.font.highlight_color = color_map[result["label"]]
    
    # Save to BytesIO
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

def _calculate_statistics(sentences: List[Dict[str, Any]], results: List[Dict[str, Any]]) -> Dict[str, Any]:
    """Calculate content statistics"""
    char_counts = {"info": 0, "promo": 0, "risk": 0}
    
    for result in results:
        idx = result["idx"]
        sentence = sentences[idx]["content"]
        
        if "spans" in result:
            for span in result["spans"]:
                char_count = span["end"] - span["start"]
                char_counts[span["label"]] += char_count
        else:
            char_counts[result["label"]] += len(sentence)
    
    total_chars = sum(char_counts.values())
    
    return {
        'info_count': char_counts["info"],
        'promo_count': char_counts["promo"],
        'risk_count': char_counts["risk"],
        'info_pct': round((char_counts["info"] / total_chars) * 100, 1) if total_chars > 0 else 0,
        'promo_pct': round((char_counts["promo"] / total_chars) * 100, 1) if total_chars > 0 else 0,
        'risk_pct': round((char_counts["risk"] / total_chars) * 100, 1) if total_chars > 0 else 0,
        'total_items': len(results)
    }

def _get_rtf_color_code(label: str) -> str:
    """Get RTF color code for classification label"""
    color_codes = {
        'info': '1',    # Light blue
        'promo': '2',   # Light coral
        'risk': '3'     # Light green
    }
    return color_codes.get(label, '4')  # Default to black

def _rtf_escape(text: str) -> str:
    """Escape text for RTF format"""
    # Escape RTF special characters
    text = text.replace('\\', '\\\\')
    text = text.replace('{', '\\{')
    text = text.replace('}', '\\}')
    
    # Handle Unicode characters
    result = []
    for char in text:
        if ord(char) > 127:
            result.append(f'\\u{ord(char)}?')
        else:
            result.append(char)
    
    return ''.join(result)

def _generate_rtf_with_structure(sentences: List[Dict[str, Any]], results: List[Dict[str, Any]], 
                               webpage_data: Dict[str, Any]) -> str:
    """Generate RTF content preserving webpage structure"""
    from bs4 import BeautifulSoup
    
    structure_html = webpage_data.get('structure', '')
    if not structure_html:
        return ""
    
    # Build classification lookup
    classification_map = _build_classification_map(sentences, results)
    
    # Parse structure
    soup = BeautifulSoup(structure_html, 'html.parser')
    
    # Convert HTML structure to RTF while preserving layout
    rtf_parts = []
    _convert_html_to_rtf(soup, classification_map, rtf_parts)
    
    return "".join(rtf_parts)

def _generate_html_with_structure(sentences: List[Dict[str, Any]], results: List[Dict[str, Any]], 
                                 webpage_data: Dict[str, Any]) -> str:
    """Generate HTML content preserving webpage structure"""
    from bs4 import BeautifulSoup
    
    structure_html = webpage_data.get('structure', '')
    if not structure_html:
        return ""
    
    # Build classification lookup
    classification_map = _build_classification_map(sentences, results)
    
    # Parse and apply classifications to structure
    soup = BeautifulSoup(structure_html, 'html.parser')
    _apply_classifications_to_dom(soup, classification_map)
    
    return str(soup)

def _generate_docx_with_structure(doc, sentences: List[Dict[str, Any]], results: List[Dict[str, Any]], 
                                 webpage_data: Dict[str, Any]):
    """Generate DOCX content preserving webpage structure"""
    from bs4 import BeautifulSoup
    from docx.enum.text import WD_COLOR_INDEX
    
    structure_html = webpage_data.get('structure', '')
    if not structure_html:
        return
    
    # Build classification lookup
    classification_map = _build_classification_map(sentences, results)
    
    # Parse structure
    soup = BeautifulSoup(structure_html, 'html.parser')
    
    # Color mapping
    color_map = {
        'info': WD_COLOR_INDEX.TURQUOISE,
        'promo': WD_COLOR_INDEX.PINK, 
        'risk': WD_COLOR_INDEX.BRIGHT_GREEN
    }
    
    # Convert HTML elements to Word elements
    _convert_html_to_docx(soup, doc, classification_map, color_map)

def _build_classification_map(sentences: List[Dict[str, Any]], 
                            results: List[Dict[str, Any]]) -> Dict[str, Any]:
    """Build a lookup map for applying classifications"""
    classification_map = {}
    
    for result in results:
        idx = result["idx"]
        sentence = sentences[idx]["content"]
        
        # Store both the original text and its classification
        classification_map[sentence] = result
        classification_map[sentence.lower()] = result
        
        # Also store sentence fragments for partial matching
        if len(sentence) > 50:
            words = sentence.split()
            if len(words) > 3:
                # Create fragment keys for better matching
                start_fragment = ' '.join(words[:3])
                end_fragment = ' '.join(words[-3:])
                classification_map[start_fragment] = result
                classification_map[end_fragment] = result
    
    return classification_map

def _apply_classifications_to_dom(element, classification_map: Dict[str, Any]):
    """Walk through DOM elements and apply classifications (same as in rendering.py)"""
    from bs4 import BeautifulSoup, NavigableString
    
    color_map = {"info": "lightblue", "promo": "lightcoral", "risk": "lightgreen"}
    
    if isinstance(element, NavigableString):
        return
    
    # Process text nodes
    for child in list(element.children):
        if isinstance(child, NavigableString):
            text_content = str(child).strip()
            if text_content and len(text_content) > 10:  # Only process substantial text
                # Try to find classification for this text
                result = _find_text_classification(text_content, classification_map)
                
                if result:
                    # Apply classification
                    if "spans" in result:
                        # Use phrase-level classification
                        classified_html = _apply_spans_to_text(text_content, result["spans"], color_map)
                    else:
                        # Use sentence-level classification
                        color = color_map.get(result["label"], "lightgray")
                        escaped_text = html.escape(text_content)
                        classified_html = f'<span style="background-color: {color};">{escaped_text}</span>'
                    
                    # Replace text with classified version
                    new_soup = BeautifulSoup(classified_html, 'html.parser')
                    child.replace_with(new_soup)
        else:
            # Recursively process child elements
            _apply_classifications_to_dom(child, classification_map)

def _find_text_classification(text: str, classification_map: Dict[str, Any]) -> Optional[Dict[str, Any]]:
    """Find the best classification match for a piece of text"""
    # Try exact match first
    result = classification_map.get(text) or classification_map.get(text.lower())
    if result:
        return result
    
    # Try partial matching for longer texts
    if len(text) > 30:
        words = text.lower().split()
        if len(words) > 3:
            # Try matching start and end fragments
            start_fragment = ' '.join(words[:3])
            end_fragment = ' '.join(words[-3:])
            
            result = classification_map.get(start_fragment) or classification_map.get(end_fragment)
            if result:
                return result
    
    # Try substring matching (less precise but catches more cases)
    text_lower = text.lower()
    for key, result in classification_map.items():
        if isinstance(key, str) and len(key) > 20:
            if key.lower() in text_lower or text_lower in key.lower():
                return result
    
    return None

def _apply_spans_to_text(text: str, spans: List[Dict[str, Any]], 
                        color_map: Dict[str, str]) -> str:
    """Apply phrase-level span classifications to text"""
    if not spans:
        return html.escape(text)
    
    # Sort spans by start position
    sorted_spans = sorted(spans, key=lambda x: x['start'])
    
    result_html = ""
    for span in sorted_spans:
        start, end, label = span['start'], span['end'], span['label']
        
        # Ensure bounds are valid for current text
        if start >= len(text) or end > len(text) or start >= end:
            continue
            
        span_text = text[start:end]
        color = color_map.get(label, "lightgray")
        escaped_text = html.escape(span_text)
        result_html += f'<span style="background-color: {color};">{escaped_text}</span>'
    
    return result_html if result_html else html.escape(text)

def _convert_html_to_rtf(element, classification_map: Dict[str, Any], rtf_parts: List[str]):
    """Convert HTML elements to RTF format while preserving structure"""
    from bs4 import NavigableString
    
    if isinstance(element, NavigableString):
        text_content = str(element).strip()
        if text_content:
            # Try to find classification for this text
            result = _find_text_classification(text_content, classification_map)
            
            if result:
                if "spans" in result:
                    # Handle phrase-level spans
                    for span in result["spans"]:
                        text_part = text_content[span["start"]:span["end"]]
                        color_code = _get_rtf_color_code(span["label"])
                        escaped_text = _rtf_escape(text_part)
                        rtf_parts.append(f"{{\\highlight{color_code} {escaped_text}}}")
                else:
                    # Handle sentence-level classification
                    color_code = _get_rtf_color_code(result["label"])
                    escaped_text = _rtf_escape(text_content)
                    rtf_parts.append(f"{{\\highlight{color_code} {escaped_text}}}")
            else:
                # No classification, add as plain text
                rtf_parts.append(_rtf_escape(text_content))
        return
    
    # Handle HTML elements
    tag_name = element.name.lower() if element.name else ""
    
    # Add appropriate RTF formatting based on HTML tag
    if tag_name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
        rtf_parts.append("\\par{\\b ")
        for child in element.children:
            _convert_html_to_rtf(child, classification_map, rtf_parts)
        rtf_parts.append("}\\par")
    elif tag_name == 'p':
        rtf_parts.append("\\par")
        for child in element.children:
            _convert_html_to_rtf(child, classification_map, rtf_parts)
        rtf_parts.append("\\par")
    elif tag_name in ['div', 'section', 'article']:
        rtf_parts.append("\\par")
        for child in element.children:
            _convert_html_to_rtf(child, classification_map, rtf_parts)
        rtf_parts.append("\\par")
    elif tag_name in ['ul', 'ol']:
        rtf_parts.append("\\par")
        for child in element.children:
            _convert_html_to_rtf(child, classification_map, rtf_parts)
        rtf_parts.append("\\par")
    elif tag_name == 'li':
        rtf_parts.append("\\par• ")
        for child in element.children:
            _convert_html_to_rtf(child, classification_map, rtf_parts)
    else:
        # For other elements, just process children
        for child in element.children:
            _convert_html_to_rtf(child, classification_map, rtf_parts)

def _convert_html_to_docx(element, doc, classification_map: Dict[str, Any], color_map: Dict[str, Any]):
    """Convert HTML elements to DOCX format while preserving structure"""
    from bs4 import NavigableString
    from docx.enum.text import WD_COLOR_INDEX
    
    if isinstance(element, NavigableString):
        return
    
    tag_name = element.name.lower() if element.name else ""
    
    # Handle different HTML elements
    if tag_name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
        level = int(tag_name[1])
        heading_text = element.get_text().strip()
        if heading_text:
            doc.add_heading(heading_text, level=min(level, 3))
    elif tag_name == 'p':
        para_text = element.get_text().strip()
        if para_text:
            para = doc.add_paragraph()
            _add_classified_text_to_paragraph(para, para_text, classification_map, color_map)
    elif tag_name in ['div', 'section', 'article']:
        # Process children elements
        for child in element.children:
            _convert_html_to_docx(child, doc, classification_map, color_map)
    elif tag_name in ['ul', 'ol']:
        # Handle lists
        for li in element.find_all('li', recursive=False):
            li_text = li.get_text().strip()
            if li_text:
                para = doc.add_paragraph(style='List Bullet' if tag_name == 'ul' else 'List Number')
                _add_classified_text_to_paragraph(para, li_text, classification_map, color_map)
    else:
        # For other elements, process children
        for child in element.children:
            _convert_html_to_docx(child, doc, classification_map, color_map)

def _add_classified_text_to_paragraph(paragraph, text: str, classification_map: Dict[str, Any], color_map: Dict[str, Any]):
    """Add classified text to a Word paragraph with highlighting"""
    # Try to find classification for this text
    result = _find_text_classification(text, classification_map)
    
    if result:
        if "spans" in result:
            # Handle phrase-level spans
            for span in result["spans"]:
                text_part = text[span["start"]:span["end"]]
                run = paragraph.add_run(text_part)
                run.font.highlight_color = color_map[span["label"]]
        else:
            # Handle sentence-level classification
            run = paragraph.add_run(text)
            run.font.highlight_color = color_map[result["label"]]
    else:
        # No classification, add as plain text
        paragraph.add_run(text)

def get_google_docs_import_instructions() -> str:
    """Return instructions for importing into Google Docs"""
    return """
# Google Docs Import Instructions

## Option 1: RTF File (Recommended for colors)
1. Download the .rtf file
2. Go to Google Docs (docs.google.com)
3. Click "Blank" to create new document
4. Go to File > Open
5. Click "Upload" tab and select your .rtf file
6. Colors and formatting will be preserved!

## Option 2: HTML File  
1. Download the .html file
2. Open it in any web browser
3. Select all content (Ctrl+A)
4. Copy (Ctrl+C)
5. Paste into Google Docs
6. Formatting should transfer

## Option 3: DOCX File (if available)
1. Download the .docx file  
2. Upload directly to Google Docs
3. Google will automatically convert it

## Pro Tips:
- RTF format preserves colors best
- After import, you can edit normally in Google Docs
- Use "Print" from Google Docs for high-quality PDF export
"""